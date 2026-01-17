import streamlit as st
import pdfplumber
from docx import Document
from docx.shared import Pt
from io import BytesIO
import re
import os
os.environ["STREAMLIT_WATCHER_TYPE"] = "none"

# --- CONTACT INFO EXTRACTION (Regex + Heuristics) ---
def extract_contact_info(text):
    email_match = re.search(r'[\w\.-]+@[\w\.-]+', text)
    email = email_match.group(0) if email_match else "Unknown"

    phone_match = re.search(r'(\+?\d{1,3}[\s-]?)?\(?\d{3}\)?[\s-]?\d{3}[\s-]?\d{4}', text)
    phone = phone_match.group(0) if phone_match else "Unknown"

    lines = [l.strip() for l in text.split("\n") if l.strip()]
    name = lines[0] if lines else "Unknown"

    return {"name": name, "email": email, "phone": phone}

# --- CLEANING FOR ATS ---
def clean_text_for_ats(text):
    date_pattern = r'([A-Za-z]{3,9}|\d{1,2})[\s,./-]{1,2}(\d{2,4})'
    def date_fix(match):
        return f"{match.group(1)}/{match.group(2)}"
    return re.sub(date_pattern, date_fix, text)

# --- DOCX GENERATOR ---
def create_ats_docx(parsed_data, raw_text, ats_type="workday"):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    header = doc.add_paragraph()
    run = header.add_run(f"{parsed_data['name']}\n{parsed_data['email']} | {parsed_data['phone']}")
    run.bold = True
    header.paragraph_format.space_after = Pt(12)

    if ats_type == "workday":
        headers = ["SUMMARY", "EXPERIENCE", "EDUCATION", "SKILLS", "PROJECTS", "CERTIFICATIONS"]
    elif ats_type == "lever":
        headers = ["SUMMARY", "WORK EXPERIENCE", "EDUCATION", "KEY SKILLS", "PROJECTS"]
    elif ats_type == "ashby":
        headers = ["SUMMARY", "BACKGROUND", "EDUCATION", "SKILLS", "PROJECTS"]
    else:
        headers = ["SUMMARY", "EXPERIENCE", "EDUCATION", "SKILLS"]

    for line in raw_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        is_header = any(h in line.upper() for h in headers) and len(line) < 40
        if is_header:
            para = doc.add_paragraph()
            run = para.add_run(line.upper())
            run.bold = True
            para.paragraph_format.space_before = Pt(12)
        else:
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- STREAMLIT UI ---
st.set_page_config(page_title="ApplyBolt — ATS Resume Converter", page_icon="⚡", layout="centered")

# Custom CSS
st.markdown("""
    <style>
    body { background-color: #f0f2f6; font-family: 'Segoe UI', sans-serif; }
    h1 { text-align: center; color: #0078d4; font-weight: 700; }
    h2, h3 { color: #333333; }
    .stButton>button {
        width: 100%; border-radius: 8px; height: 3em;
        background: linear-gradient(90deg, #0078d4, #005a9e);
        color: white; font-weight: bold; border: none;
    }
    .stDownloadButton>button {
        width: 100%; border-radius: 8px; height: 3em;
        background: linear-gradient(90deg, #28a745, #218838);
        color: white; font-weight: bold; border: none;
    }
    .footer { text-align: center; margin-top: 50px; font-size: 0.9em; color: #666666; }
    </style>
""", unsafe_allow_html=True)

# Header
st.markdown("<h1>⚡ ApplyBolt — ATS Resume Converter ⚡</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align:center;color:#555;'>Upload once → Get Workday, Lever, AshbyHQ-ready resumes instantly.</p>", unsafe_allow_html=True)

# Payment Section
st.markdown("### 💳 Payment Options")
st.markdown("#### PayPal (USD)")
st.markdown("""
    <form action="https://www.paypal.com/donate" method="post" target="_top">
    <input type="hidden" name="business" value="your-paypal-email@example.com" />
    <input type="hidden" name="currency_code" value="USD" />
    <input type="hidden" name="amount" value="5.00" />
    <input type="submit" value="Pay $5 via PayPal" style="background:#0078d4;color:white;border:none;padding:10px 20px;border-radius:8px;font-weight:bold;" />
    </form>
""", unsafe_allow_html=True)

st.markdown("#### UPI (India)")
st.image("upi_qr.png", caption="Scan to Pay via UPI")  # Add your QR image in repo
st.markdown("<p style='text-align:center;'>UPI ID: yourname@upi</p>", unsafe_allow_html=True)

# Resume Upload
uploaded_file = st.file_uploader("📂 Upload Resume (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
ats_choice = st.selectbox("Choose ATS Format", ["Workday", "Lever", "AshbyHQ"])

if uploaded_file:
    with st.spinner("Analyzing resume..."):
        ext = uploaded_file.name.split('.')[-1].lower()
        if ext == "pdf":
            with pdfplumber.open(uploaded_file) as pdf:
                raw_text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
        elif ext == "docx":
            doc = Document(uploaded_file)
            raw_text = "\n".join([p.text for p in doc.paragraphs])
        else:
            raw_text = uploaded_file.read().decode("utf-8")

        parsed_data = extract_contact_info(raw_text)
        cleaned_text = clean_text_for_ats(raw_text)
        final_docx = create_ats_docx(parsed_data, cleaned_text, ats_choice.lower())

        st.success(f"✅ Resume optimized for {ats_choice}!")
        st.download_button(
            label=f"⬇️ Download {ats_choice}-Ready Resume (.docx)",
            data=final_docx,
            file_name=f"{ats_choice}_Ready_{uploaded_file.name.split('.')[0]}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

st.markdown("<div class='footer'>Made with ⚡ ApplyBolt | © 2026</div>", unsafe_allow_html=True)
